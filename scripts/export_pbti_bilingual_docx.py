from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
APP_JS = ROOT / "PBTI" / "app.js"
LOCALIZATION_JS = ROOT / "PBTI" / "localization.js"
OUTPUT = ROOT / "output" / "PBTI咖啡熊格测试中英文对照文案-优化完成.docx"

# compact_reference_guide preset, with a restrained coffee-colored named override.
INK = "2D2118"
COFFEE = "70451F"
ACCENT = "B98A55"
MUTED = "7F6B58"
PALE = "F5EADC"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
ASCII_FONT = "Calibri"
CJK_FONT = "Fusion Pixel"

TYPE_ORDER = ["MASTER", "NEED", "YUMMY", "NEW", "PHOTO", "REPORT", "DESSERT", "BREW"]

DIMENSION_LEVELS_EN = {
    "D1": {
        "L": "If it is drinkable, that is enough—no coffee essay required.",
        "M": "You can tell whether a cup is good and occasionally ask a question or two.",
        "H": "Origin, roast, and processing method all register automatically.",
    },
    "D2": {
        "L": "You can function perfectly well without coffee.",
        "M": "Coffee makes the day steadier, but you will not shut down without it.",
        "H": "Until the coffee arrives, the day has not officially booted up.",
    },
    "D3": {
        "L": "You do not need milk, sugar, or toppings to rescue the cup.",
        "M": "It depends on the day: straight sometimes, with extras at other times.",
        "H": "Milky, fruity, or chocolatey—the smoother it tastes, the happier you are.",
    },
    "D4": {
        "L": "There are plenty of new drinks, but old favorites are the most reliable.",
        "M": "If it sounds good, you will try it; if it sounds too strange, you will pass.",
        "H": "The more limited it is, the more tempting; the less familiar, the more you want it.",
    },
    "D5": {
        "L": "Drink first. The photo does not matter.",
        "M": "If it looks good, you will take a photo; if not, you will still enjoy it.",
        "H": "The cup, the light, and the table all need a quick inspection first.",
    },
    "D6": {
        "L": "You are usually the last person to hear about a new café nearby.",
        "M": "If you see one online, you save it and visit when you have time.",
        "H": "New cafés, menus, queues, and seating are practically your reporting beat.",
    },
    "D7": {
        "L": "Coffee is enough on its own; it does not need a supporting act.",
        "M": "If you are hungry, you might add something to eat.",
        "H": "Coffee is best with bread, cake, a light meal, or even a full meal beside it.",
    },
    "D8": {
        "L": "Leave it to the barista—doing it yourself is asking for trouble.",
        "M": "Drip bags and capsules are fine; complicated equipment can wait.",
        "H": "Wherever you are, you can improvise a cup of coffee your way.",
    },
}

RESULT_COPY = [
    ("你的主类型", "Your Coffee Type"),
    ("多维选项权重综合匹配", "Matched across all eight dimensions"),
    (
        "每个选择都会同时影响多个熊格倾向，雷达轮廓共同形成当前结果。",
        "Every answer influences several coffee tendencies; together, they shape your final profile.",
    ),
    ("PBTI 咖啡八维轮廓", "Your Eight-Dimension Coffee Profile"),
    ("熊格简单解读", "Your Bearsonality"),
    (
        "该测试仅供娱乐，结果只是你当下咖啡偏好的一个轻量切片，不代表固定性格或专业判断。",
        "This quiz is just for fun. Your result is a light snapshot of your current coffee preferences—not a fixed personality label or a professional assessment.",
    ),
    ("重新测试", "Retake the Quiz"),
    ("购买", "Shop"),
]

CREATOR_COPY = [
    (
        "PBTI 想记录的不是“你是哪一种咖啡”，而是你如何把咖啡放进日常生活里：续命、尝新、探店、拍照、配餐，或是在家认真给自己弄一杯。",
        "PBTI is not trying to decide which coffee you are. It looks at how coffee fits into your everyday life: keeping you going, chasing new flavors, exploring cafés, taking photos, pairing food, or making a proper cup at home.",
    ),
    (
        "八维轮廓不是评分高低，而是一个偏好形状。它会随着作息、心情、城市和最近喝到的那杯咖啡改变。",
        "The eight-dimension profile is not a scorecard; it is the shape of your preferences. It can shift with your routine, mood, city, and even the last cup you had.",
    ),
    (
        "希望这个结果能像一次轻松的咖啡聊天：让你更理解自己的小习惯，也多一个和朋友互相调侃的理由。",
        "Think of this result as a relaxed coffee chat: a chance to notice your little habits and gain one more thing to tease your friends about.",
    ),
]


def set_run_font(run, size=None, bold=None, italic=None, color=None, font=ASCII_FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_font(style, size, color, bold=False):
    style.font.name = ASCII_FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ASCII_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=8.5, color=MUTED)


def extract_chinese_data(source: str):
    dimensions = {
        key: {"name": name, "model": model}
        for key, name, model in re.findall(r"(D\d+):\{name:`([^`]*)`,model:`([^`]*)`\}", source)
    }
    explanations = {
        key: {"L": low, "M": mid, "H": high}
        for key, low, mid, high in re.findall(
            r"(D\d+):\{L:`([^`]*)`,M:`([^`]*)`,H:`([^`]*)`\}", source
        )
    }

    types = {}
    for code, actual_code, display_code, cn, intro, desc in re.findall(
        r"([A-Z]+):\{code:`([^`]*)`,(?:displayCode:`([^`]*)`,)?cn:`([^`]*)`,intro:`([^`]*)`,desc:`([^`]*)`\}",
        source,
    ):
        if code == actual_code and code in TYPE_ORDER:
            types[code] = {
                "display_code": display_code or code,
                "cn": cn,
                "intro": intro,
                "desc": desc,
            }

    questions = []
    for qid, dim, text, options_block in re.findall(
        r"\{id:`(q\d+)`,dim:`(D\d+)`,text:`([^`]*)`,options:\[(.*?)\]\}", source
    ):
        options = [
            label for label, _ in re.findall(r"\{label:`([^`]*)`,value:(\d+)\}", options_block)
        ]
        questions.append({"id": qid, "dim": dim, "text": text, "options": options})

    assert len(dimensions) == 8 and len(explanations) == 8
    assert len(types) == 8 and len(questions) == 15
    return dimensions, explanations, types, questions


def extract_english_data(source: str):
    questions = {}
    for qid, text, options_block in re.findall(
        r"\s+(q\d+): \{\s+text: '([^']*)',\s+options: \[(.*?)\s+\]\s+\}",
        source,
        re.S,
    ):
        questions[qid] = {
            "text": text,
            "options": re.findall(r"\s+'([^']*)'[,]?", options_block),
        }

    types = {}
    for code in TYPE_ORDER:
        match = re.search(
            rf"\s+{code}: \{{\s+name: '([^']*)',\s+intro: '([^']*)',\s+desc: '([^']*)'\s+\}}",
            source,
            re.S,
        )
        assert match, f"Missing English type copy: {code}"
        types[code] = {"name": match.group(1), "intro": match.group(2), "desc": match.group(3)}

    dimension_names = dict(
        re.findall(r"\s+'([^']+)': '([^']+)'[,]?", re.search(r"const DIMENSIONS = \{(.*?)\};", source, re.S).group(1))
    )
    assert len(questions) == 15 and len(types) == 8 and len(dimension_names) == 8
    return questions, types, dimension_names


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    style_font(normal, 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE_DARK, 10, 5),
    ):
        style = doc.styles[name]
        style_font(style, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bilingual_cn = doc.styles.add_style("Bilingual Chinese", WD_STYLE_TYPE.PARAGRAPH)
    bilingual_cn.base_style = normal
    style_font(bilingual_cn, 11, INK)
    bilingual_cn.paragraph_format.space_after = Pt(2)
    bilingual_cn.paragraph_format.keep_with_next = True

    bilingual_en = doc.styles.add_style("Bilingual English", WD_STYLE_TYPE.PARAGRAPH)
    bilingual_en.base_style = normal
    style_font(bilingual_en, 10.5, MUTED)
    bilingual_en.font.italic = True
    bilingual_en.paragraph_format.space_after = Pt(9)

    option_style = doc.styles.add_style("Option Pair", WD_STYLE_TYPE.PARAGRAPH)
    option_style.base_style = normal
    style_font(option_style, 10.5, INK)
    option_style.paragraph_format.left_indent = Inches(0.28)
    option_style.paragraph_format.first_line_indent = Inches(-0.28)
    option_style.paragraph_format.space_after = Pt(7)
    option_style.paragraph_format.line_spacing = 1.2

    kicker = doc.styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    kicker.base_style = normal
    style_font(kicker, 10.5, ACCENT, bold=True)
    kicker.paragraph_format.space_after = Pt(12)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("PBTI · Bilingual Copy Reference")
    set_run_font(run, 8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PBTI 咖啡熊格测试中英文对照 · ")
    set_run_font(run, 8.5, color=MUTED)
    add_page_field(footer)


def add_title_page(doc: Document):
    # editorial_cover header pattern: generous whitespace and centered title stack.
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PBTI COPY REFERENCE")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("PBTI 咖啡熊格测试")
    set_run_font(run, 29, bold=True, color=COFFEE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("中英文对照文案")
    set_run_font(run, 16, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("15 道题目 · 8 个测量维度 · 8 种熊格 · 结果页与作者说明")
    set_run_font(run, 10.5, color=MUTED)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("整理原则：熊格代码与中文名称保持原样，其他用户可见内容提供中英对照。")
    set_run_font(run, 9.5, italic=True, color=MUTED)
    doc.add_page_break()


def add_pair(doc: Document, zh: str, en: str, label: str | None = None, compact: bool = False):
    p = doc.add_paragraph(style="Bilingual Chinese")
    if label:
        r = p.add_run(f"{label}｜")
        set_run_font(r, bold=True, color=COFFEE)
    p.add_run(zh)
    if compact:
        p.paragraph_format.space_after = Pt(1)
    p = doc.add_paragraph(style="Bilingual English")
    p.add_run(en)
    if compact:
        p.paragraph_format.space_after = Pt(4)


def add_overview(doc: Document):
    doc.add_heading("文档说明｜About This Document", level=1)
    add_pair(
        doc,
        "本文档依据当前 PBTI 网页文案整理，按“中文原文在前、英文版本在后”的方式呈现。",
        "This document reflects the current PBTI website copy, with the original Chinese followed by its English version.",
    )
    add_pair(
        doc,
        "MASTER、NEED、YUMMY、NEW、CAMERA、REPORT、EAT、HOME-MASTER 为熊格代码；煮理人、咖啡续命者、咖界金舌头等中文名称属于角色正式名称，在中英文页面中均保持不变。",
        "MASTER, NEED, YUMMY, NEW, CAMERA, REPORT, EAT, and HOME-MASTER are profile codes. Official Chinese character names such as 煮理人, 咖啡续命者, and 咖界金舌头 remain unchanged in both language versions.",
    )
    doc.add_heading("内容结构｜Contents", level=2)
    for title in (
        "第一部分｜熊格简介与解读 / Profile Introductions and Interpretations",
        "第二部分｜咖啡八维解读 / Eight Coffee Dimensions",
        "第三部分｜测试题与选项 / Questions and Options",
        "第四部分｜结果页与作者说明 / Results and Creator’s Note",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title)
        set_run_font(r, 10.5, bold=True, color=COFFEE)
    doc.add_page_break()


def add_types(doc: Document, zh_types, en_types):
    doc.add_heading("第一部分｜熊格简介与解读", level=1)
    p = doc.add_paragraph(style="Bilingual English")
    p.add_run("Part 1 · Profile Introductions and Interpretations")
    for index, code in enumerate(TYPE_ORDER):
        data_zh = zh_types[code]
        data_en = en_types[code]
        if index in (2, 4, 6):
            doc.add_page_break()
        doc.add_heading(f"{data_zh['display_code']} · {data_zh['cn']}", level=2)
        add_pair(doc, data_zh["intro"], data_en["intro"], "简介 / Intro")
        add_pair(doc, data_zh["desc"], data_en["desc"], "解读 / Interpretation")
    doc.add_page_break()


def add_dimensions(doc: Document, dimensions, explanations, dimension_names_en):
    doc.add_heading("第二部分｜咖啡八维解读", level=1)
    p = doc.add_paragraph(style="Bilingual English")
    p.add_run("Part 2 · Eight Coffee Dimensions")
    for index, key in enumerate(sorted(dimensions, key=lambda value: int(value[1:]))):
        if index == 4:
            doc.add_page_break()
        zh_name = re.sub(r"^D\d+\s*", "", dimensions[key]["name"])
        en_name = dimension_names_en[zh_name]
        doc.add_heading(f"{key} · {zh_name} / {en_name}", level=2)
        for level, zh_level in (("L", "低"), ("M", "中"), ("H", "高")):
            add_pair(doc, explanations[key][level], DIMENSION_LEVELS_EN[key][level], f"{level} · {zh_level}")
    doc.add_page_break()


def add_questions(doc: Document, questions_zh, questions_en, dimensions, dimension_names_en):
    doc.add_heading("第三部分｜测试题与选项", level=1)
    p = doc.add_paragraph(style="Bilingual English")
    p.add_run("Part 3 · Questions and Options")
    for index, question in enumerate(questions_zh, 1):
        if index in (5, 9, 13):
            doc.add_page_break()
        dim_zh = re.sub(r"^D\d+\s*", "", dimensions[question["dim"]]["name"])
        dim_en = dimension_names_en[dim_zh]
        en = questions_en[question["id"]]
        doc.add_heading(f"{index:02d} · {dim_zh} / {dim_en}", level=2)
        add_pair(doc, question["text"], en["text"])
        for option_index, (zh_option, en_option) in enumerate(zip(question["options"], en["options"])):
            letter = chr(ord("A") + option_index)
            p = doc.add_paragraph(style="Option Pair")
            r = p.add_run(f"{letter}  ")
            set_run_font(r, bold=True, color=COFFEE)
            p.add_run(f"{zh_option}\n")
            r = p.add_run(en_option)
            set_run_font(r, italic=True, color=MUTED)
    doc.add_page_break()


def add_results(doc: Document):
    doc.add_heading("第四部分｜结果页与作者说明", level=1)
    p = doc.add_paragraph(style="Bilingual English")
    p.add_run("Part 4 · Results and Creator’s Note")
    doc.add_heading("结果页固定文案｜Results Page Copy", level=2)
    for zh, en in RESULT_COPY:
        add_pair(doc, zh, en, compact=True)
    doc.add_heading("作者的话｜A Note from the Creator", level=2)
    for zh, en in CREATOR_COPY:
        add_pair(doc, zh, en, compact=True)


def main():
    app_source = APP_JS.read_text(encoding="utf-8")
    localization_source = LOCALIZATION_JS.read_text(encoding="utf-8")
    dimensions, explanations, zh_types, zh_questions = extract_chinese_data(app_source)
    en_questions, en_types, dimension_names_en = extract_english_data(localization_source)

    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_overview(doc)
    add_types(doc, zh_types, en_types)
    add_dimensions(doc, dimensions, explanations, dimension_names_en)
    add_questions(doc, zh_questions, en_questions, dimensions, dimension_names_en)
    add_results(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")
    print(f"Types: {len(zh_types)}; dimensions: {len(dimensions)}; questions: {len(zh_questions)}")


if __name__ == "__main__":
    main()
