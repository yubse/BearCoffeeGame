from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "PBTI" / "index.html"
OUTPUT = ROOT / "output" / "PBTI熊格测试文案全集.docx"

COFFEE = "8A5A36"
COFFEE_DARK = "3C281C"
CREAM = "F8F2EA"
CREAM_DARK = "E7D3BF"
MUTED = "7F7064"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_run_font(run, size=None, bold=None, color=None, east_asia="Arial Unicode MS"):
    run.font.name = "Arial Unicode MS"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def extract_data(html: str):
    dimensions = {
        key: {"name": name, "model": model}
        for key, name, model in re.findall(
            r"(D\d+):\{name:`([^`]*)`,model:`([^`]*)`\}", html
        )
    }
    explanations = {
        key: {"L": low, "M": mid, "H": high}
        for key, low, mid, high in re.findall(
            r"(D\d+):\{L:`([^`]*)`,M:`([^`]*)`,H:`([^`]*)`\}", html
        )
    }
    types = {}
    for code, code2, display_code, cn, intro, desc in re.findall(
        r"([A-Z]+):\{code:`([^`]*)`,(?:displayCode:`([^`]*)`,)?cn:`([^`]*)`,intro:`([^`]*)`,desc:`([^`]*)`\}",
        html,
    ):
        if code == code2 and code in {
            "MASTER", "NEED", "YUMMY", "NEW", "PHOTO", "REPORT", "DESSERT", "BREW"
        }:
            types[code] = {
                "display_code": display_code or code2,
                "cn": cn,
                "intro": intro,
                "desc": desc,
            }

    patterns = dict(re.findall(r"\{code:`([A-Z]+)`,pattern:`([^`]*)`\}", html))
    questions = []
    for qid, dim, text, options_block in re.findall(
        r"\{id:`(q\d+)`,dim:`(D\d+)`,text:`([^`]*)`,options:\[(.*?)\]\}",
        html,
    ):
        options = [
            {"label": label, "value": int(value)}
            for label, value in re.findall(
                r"\{label:`([^`]*)`,value:(\d+)\}", options_block
            )
        ]
        questions.append({"id": qid, "dim": dim, "text": text, "options": options})

    order = ["MASTER", "NEED", "YUMMY", "NEW", "PHOTO", "REPORT", "DESSERT", "BREW"]
    assert len(types) == 8, f"Expected 8 types, got {len(types)}"
    assert len(dimensions) == 8 and len(explanations) == 8
    assert len(questions) == 15
    return dimensions, explanations, [(c, types[c], patterns[c]) for c in order], questions


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COFFEE_DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 31, COFFEE_DARK, 0, 12),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 19, COFFEE, 18, 8),
        ("Heading 2", 14, COFFEE_DARK, 13, 6),
        ("Heading 3", 11.5, COFFEE, 9, 4),
    ):
        st = doc.styles[name]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st.font.size = Pt(size)
        st.font.bold = name != "Subtitle"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    quote = doc.styles.add_style("Intro Callout", WD_STYLE_TYPE.PARAGRAPH)
    quote.base_style = doc.styles["Normal"]
    quote.font.name = "Arial Unicode MS"
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    quote.font.size = Pt(12)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(COFFEE)
    quote.paragraph_format.left_indent = Inches(0.18)
    quote.paragraph_format.right_indent = Inches(0.18)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(10)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("PBTI 咖啡熊格测试 · 文案资料库")
    set_run_font(run, 8.5, color=MUTED)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PBTI · 当前网页版本全量整理")
    set_run_font(run, 8, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PBTI 文案资料库")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("咖啡熊格测试 · 全量文案整理")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("8 种熊格｜8 个维度｜15 道测试题｜结果页与界面文案")
    set_run_font(r, 10.5, color=COFFEE)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"依据 PBTI 当前网页代码整理\n生成日期：{date.today().isoformat()}")
    set_run_font(r, 9, color=MUTED)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("文档结构", level=1)
    items = [
        ("第一部分", "熊格类型文案", "按 8 种熊格分别汇总英文代码、中文名称、匹配模式、Intro 与 Desc。"),
        ("第二部分", "维度文案", "汇总 8 个测量维度，以及低、中、高三个等级的解读。"),
        ("第三部分", "测试题文案", "汇总 15 道题及全部选项，并标记所属维度。"),
        ("第四部分", "结果逻辑文案", "汇总主类型、匹配度、特殊结果与结果报告中的固定文字。"),
        ("第五部分", "页面公共文案", "汇总首页、测试页、结果页及导航中使用的主要界面文字。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("部分", "主题", "内容")):
        cell.text = text
        shade(cell, COFFEE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 9.5, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])
    for part, title, desc in items:
        cells = table.add_row().cells
        for cell, text in zip(cells, (part, title, desc)):
            cell.text = text
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("整理口径：")
    set_run_font(r, bold=True, color=COFFEE)
    p.add_run("以网页当前可运行版本中的数据与固定界面文案为准；“熊格”沿用项目现行称呼。")
    doc.add_page_break()


def add_types(doc: Document, types) -> None:
    doc.add_heading("第一部分｜熊格类型文案", level=1)
    doc.add_paragraph("以下内容按网页展示顺序整理。匹配模式中的 H／M／L 分别代表高／中／低等级。")
    for idx, (code, data, pattern) in enumerate(types):
        if idx:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(data["display_code"])
        set_run_font(r, 27, bold=True, color=COFFEE_DARK)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(data["cn"])
        set_run_font(r, 16, bold=True, color=COFFEE)

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        labels = (("熊格代码", data["display_code"]), ("匹配模式", pattern))
        for cell, (label, value) in zip(table.rows[0].cells, labels):
            shade(cell, CREAM)
            set_cell_margins(cell, top=130, bottom=130)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(f"{label}\n")
            set_run_font(r, 8.5, bold=True, color=MUTED)
            r = para.add_run(value)
            set_run_font(r, 11, bold=True, color=COFFEE_DARK)

        doc.add_heading("Intro", level=2)
        p = doc.add_paragraph(style="Intro Callout")
        p.add_run(data["intro"])
        doc.add_heading("熊格解读（Desc）", level=2)
        doc.add_paragraph(data["desc"])
    doc.add_page_break()


def add_dimensions(doc: Document, dimensions, explanations) -> None:
    doc.add_heading("第二部分｜维度文案", level=1)
    doc.add_paragraph("每个维度对应三个等级：L（低）、M（中）、H（高）。")
    for key in sorted(dimensions, key=lambda x: int(x[1:])):
        meta = dimensions[key]
        doc.add_heading(f"{meta['name']}｜{meta['model']}", level=2)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row, level in zip(table.rows, ("L", "M", "H")):
            row.cells[0].text = {"L": "L｜低", "M": "M｜中", "H": "H｜高"}[level]
            row.cells[1].text = explanations[key][level]
            shade(row.cells[0], CREAM_DARK if level == "H" else CREAM)
            for cell in row.cells:
                set_cell_margins(cell)
            for run in row.cells[0].paragraphs[0].runs:
                set_run_font(run, 9.5, bold=True, color=COFFEE)
        doc.add_paragraph()
    doc.add_page_break()


def add_questions(doc: Document, questions, dimensions) -> None:
    doc.add_heading("第三部分｜测试题文案", level=1)
    doc.add_paragraph("选项值 1／2／3 分别进入对应维度的累计分数。")
    for index, q in enumerate(questions, 1):
        dim_label = dimensions[q["dim"]]["name"]
        doc.add_heading(f"第 {index} 题｜{dim_label}", level=2)
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        r = p.add_run(q["text"])
        set_run_font(r, 11, bold=True, color=COFFEE_DARK)
        for option_index, option in enumerate(q["options"]):
            letter = chr(ord("A") + option_index)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.add_run(f"{letter}. {option['label']}  ")
            r = p.add_run(f"〔值：{option['value']}〕")
            set_run_font(r, 8.5, color=MUTED)
    doc.add_page_break()


def add_result_copy(doc: Document) -> None:
    doc.add_heading("第四部分｜结果逻辑与结果页文案", level=1)
    sections = [
        ("常规结果", [
            "你的主类型",
            "匹配度 {相似度}% · 精准命中 {命中维度数}/8 维",
            "维度命中度较高，当前结果可视为你的第一熊格画像。",
        ]),
        ("隐藏熊格结果", [
            "隐藏熊格已激活",
            "匹配度 100% · 特殊因子已接管",
            "特殊因子过强，系统已直接跳过常规熊格审判。",
        ]),
        ("兜底熊格结果", [
            "系统强制兜底",
            "标准熊格库最高匹配仅 {相似度}%",
            "标准熊格库对你的脑回路集体罢工了，于是系统把你强制分配给了兜底熊格。",
        ]),
        ("结果报告固定标题", [
            "你的专属PBTI测试报告已生成",
            "你的测试类型",
            "类型数据",
            "熊格解读",
            "重新测试",
        ]),
    ]
    for title, lines in sections:
        doc.add_heading(title, level=2)
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("结果匹配规则说明", level=2)
    doc.add_paragraph(
        "系统将 8 个维度的结果转换为 H／M／L，再与每种熊格的匹配模式计算距离、精准命中数和相似度；"
        "常规情况下选择距离最小的熊格。若最高相似度低于 60%，且存在兜底类型，则显示兜底熊格。"
    )
    doc.add_page_break()


def add_interface_copy(doc: Document) -> None:
    doc.add_heading("第五部分｜页面公共文案", level=1)
    groups = [
        ("项目信息", ["PBTI", "咖啡熊格测试", "by Lovstudio", "8 种熊格"]),
        ("首页与开屏", ["PBTI", "咖啡熊格测试", "开始测试", "切换风格"]),
        ("主站导航", ["咖啡屋", "熊格测试", "咖啡地图", "未开放"]),
        ("测试过程", ["第 {序号} 题", "补充题", "{已完成题数} / {总题数}", "提交测试", "返回"]),
        ("结果展示", ["你的专属PBTI测试报告已生成", "你的测试类型", "类型数据", "熊格解读", "重新测试"]),
        ("站点页脚", ["Powered by XBTI", "PBTI"]),
    ]
    for title, lines in groups:
        doc.add_heading(title, level=2)
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("附注", level=2)
    doc.add_paragraph(
        "本文档重点收录面向用户的 PBTI 测试内容。网页框架自身的开发者推广语、命令行安装说明和外部购买链接，"
        "不属于 8 种熊格的测试文案，因此未并入熊格正文。"
    )


def main() -> None:
    html = SOURCE.read_text(encoding="utf-8")
    dimensions, explanations, types, questions = extract_data(html)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_types(doc, types)
    add_dimensions(doc, dimensions, explanations)
    add_questions(doc, questions, dimensions)
    add_result_copy(doc)
    add_interface_copy(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")
    print(f"Types: {len(types)}, dimensions: {len(dimensions)}, questions: {len(questions)}")


if __name__ == "__main__":
    main()
